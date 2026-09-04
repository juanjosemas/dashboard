#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Regenera dashboard.html con numeros CORRECTOS directamente del CSV y .docx
"""
import csv
import io
import re
import collections
import json
from docx import Document

def parse_euro_amount(s):
    """Parse European amount: handles '8875,66' and '116432.26' and '1.699,09'"""
    s = s.strip().replace('\u20ac', '').strip()
    # Detect format: if comma followed by 1-2 digits at end -> comma is decimal
    if re.search(r',\d{1,2}$', s):
        s = s.replace('.', '').replace(',', '.')
    elif re.search(r'\.\d{1,2}$', s):
        s = s.replace(',', '')
    else:
        s = s.replace('.', '').replace(',', '')
    try:
        return float(s)
    except:
        return 0.0

# ===== 1. CERTIFICACIONES =====
doc_cert = Document(r'C:\Users\jjmax\Downloads\1\dashboard\CERTIFICACIONES.docx')

certificaciones = []
for p in doc_cert.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    # Find ALL euro amounts in the line and take the LAST one (the total)
    matches = list(re.finditer(r'([\d.,]+)\s*\u20ac', text))
    if not matches:
        continue
    # Last match is the total amount
    last_match = matches[-1]
    importe = parse_euro_amount(last_match.group(1))
    
    # Project name = everything before the dashes that precede the final number
    proyecto = text[:last_match.start()].strip()
    # Remove trailing dashes and spaces
    proyecto = re.sub(r'[\s\-]+$', '', proyecto).strip()
    # Remove formula parts like "(3288,00 + 63277,34) ="
    proyecto = re.sub(r'\([\d.,\s\+\-]+\)\s*=\s*$', '', proyecto).strip()
    proyecto = re.sub(r'[\s\-]+$', '', proyecto).strip()
    
    if importe > 0 and proyecto:
        certificaciones.append({'nombre': proyecto, 'importe': importe})

total_cert = sum(c['importe'] for c in certificaciones)
print("Certificaciones: %d proyectos, total %s EUR" % (len(certificaciones), "{:,.2f}".format(total_cert)))
for c in certificaciones:
    pct = c['importe'] / total_cert * 100
    print("  %-60s %s EUR (%.2f%%)" % (c['nombre'], "{:>12,.2f}".format(c['importe']), pct))

# ===== 2. MANO DE OBRA (2025 + 2026) =====
doc_mo = Document(r'C:\Users\jjmax\Downloads\1\dashboard\GASTOS MANO DE OBRA.docx')

mano_obra_all = []  # ALL years combined (like the original dashboard)
current_year = 2026

for p in doc_mo.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    if '2025' in text:
        current_year = 2025
        continue
    if '2026' in text:
        current_year = 2026
        continue
    if text.startswith('GASTOS'):
        continue
    
    # Normalize: replace non-ASCII chars with spaces to fix encoding artifacts
    clean = re.sub(r'[^\x00-\x7f]+', ' ', text)
    clean = re.sub(r'\s+', ' ', clean).strip()
    
    # Format 1: PROYECTO --- XXXX HORAS A 20E - TOTAL XXXX
    # Find 'HORAS A' pattern and work backwards to extract name and hours
    horas_match = re.search(r'(\d+)\s*HORAS?\s*A\s*(\d+).*?TOTAL\s*([\d.,]+)', clean, re.IGNORECASE)
    if horas_match:
        horas_val = int(horas_match.group(1))
        tarifa_val = int(horas_match.group(2))
        coste_val = parse_euro_amount(horas_match.group(3))
        # Extract name: everything before the hours number
        name_end = horas_match.start()
        proyecto = clean[:name_end].strip()
        # Remove trailing dashes/spaces from name
        proyecto = re.sub(r'[\s\-]+$', '', proyecto).strip()
        if proyecto:
            mano_obra_all.append({'proyecto': proyecto, 'horas': horas_val, 'tarifa': tarifa_val, 'coste': coste_val, 'year': current_year})
            continue
    
    # Format 2: PROYECTO --- TOTAL XXXX,XX (no hours specified)
    match2 = re.search(r'(.+?)[\s\-]+TOTAL\s*([\d.,]+)', clean, re.IGNORECASE)
    if match2:
        proyecto = match2.group(1).strip()
        coste = parse_euro_amount(match2.group(2))
        horas = 0  # No hours specified - use coste directly
        mano_obra_all.append({'proyecto': proyecto, 'horas': horas, 'tarifa': 0, 'coste': coste, 'year': current_year})

total_horas_mo = sum(m['horas'] for m in mano_obra_all)
def mo_cost(m):
    """Get MO cost: use stated total for entries without tarifa (MURO VECINO, OBRA CAMPO)."""
    if m['tarifa'] > 0:
        return m['horas'] * m['tarifa']
    return m['coste']  # Use the exact stated total
total_coste_mo = sum(mo_cost(m) for m in mano_obra_all)
print("\nMano de Obra TOTAL: %d entradas, %d horas, %s EUR" % (len(mano_obra_all), total_horas_mo, "{:,.2f}".format(total_coste_mo)))
for m in mano_obra_all:
    print("  %-50s %5d h x %dEUR = %sEUR  [%d]" % (m['proyecto'], m['horas'], m['tarifa'], "{:>10,.2f}".format(m['coste']), m['year']))

# ===== 3. CSV GASTOS =====
csv_path = r'C:\Users\jjmax\Downloads\1\dashboard\ECO_STRUCT_-_Workspace_Gastos.csv'

with open(csv_path, 'r', encoding='latin-1') as f:
    content = f.read()

reader = csv.DictReader(io.StringIO(content), delimiter=';')

gg_total = 0.0
gg_count = 0
veh_total = 0.0
veh_count = 0
directos_por_proyecto = collections.defaultdict(lambda: {'total': 0.0, 'count': 0, 'facturas': []})

for row in reader:
    proyecto = ''
    importe_str = '0'
    titulo = ''
    proveedor = ''
    fecha = ''
    codigo = ''
    estado = ''
    
    for key, val in row.items():
        if key is None:
            continue
        key_lower = key.lower()
        val = val.strip() if val else ''
        if 'proyecto' == key_lower:
            proyecto = val.upper()
        elif 'importe' == key_lower:
            importe_str = val
        elif 't' in key_lower and 'tulo' in key_lower:
            titulo = val
        elif 'proveedor' in key_lower and 'id' not in key_lower:
            proveedor = val
        elif 'fecha' in key_lower and 'imputa' in key_lower:
            fecha = val
        elif 'digo' in key_lower or 'odigo' in key_lower:
            codigo = val
        elif 'estado' == key_lower:
            estado = val
    
    importe = parse_euro_amount(importe_str)
    
    if 'GASTOS GENERALES' in proyecto:
        gg_total += importe
        gg_count += 1
    elif 'VEHICULOS' in proyecto or 'VEH\u00cdCULOS' in proyecto:
        veh_total += importe
        veh_count += 1
    elif proyecto:
        directos_por_proyecto[proyecto]['total'] += importe
        directos_por_proyecto[proyecto]['count'] += 1
        directos_por_proyecto[proyecto]['facturas'].append({
            'codigo': codigo, 'fecha': fecha, 'titulo': titulo,
            'importe': importe, 'proveedor': proveedor, 'estado': estado
        })

total_gastos_comunes = gg_total + veh_total
total_directos_csv = sum(d['total'] for d in directos_por_proyecto.values())
total_facturas_dir = sum(d['count'] for d in directos_por_proyecto.values())

# Collect ALL individual direct invoices as flat list for the 'Facturas por Obra' tab
all_facturas_dir = []
for proj_csv, d in directos_por_proyecto.items():
    for f in d['facturas']:
        all_facturas_dir.append({
            'proyecto': proj_csv,
            'codigo': f.get('codigo', ''),
            'fecha': f.get('fecha', ''),
            'titulo': f.get('titulo', ''),
            'proveedor': f.get('proveedor', ''),
            'estado': f.get('estado', ''),
            'importe': f.get('importe', 0)
        })
# Sort by project then by amount descending
all_facturas_dir.sort(key=lambda x: (x['proyecto'], -abs(x['importe'])))

print("\nGastos Generales: %d facturas, %s EUR" % (gg_count, "{:,.2f}".format(gg_total)))
print("Vehiculos: %d facturas, %s EUR" % (veh_count, "{:,.2f}".format(veh_total)))
print("Gastos Comunes: %s EUR" % "{:,.2f}".format(total_gastos_comunes))
print("Gastos Directos: %d facturas, %s EUR" % (total_facturas_dir, "{:,.2f}".format(total_directos_csv)))
for proj in sorted(directos_por_proyecto.keys(), key=lambda x: directos_por_proyecto[x]['total'], reverse=True):
    d = directos_por_proyecto[proj]
    print("  %-65s %s EUR (%d facturas)" % (proj[:65], "{:>12,.2f}".format(d['total']), d['count']))

# ===== 4. MAPEO CSV -> CERTIFICACIONES =====
csv_to_cert = {}
for proj_csv in directos_por_proyecto:
    proj_upper = proj_csv.upper()
    for cert in certificaciones:
        cert_upper = cert['nombre'].upper()
        if cert_upper in proj_upper or proj_upper in cert_upper:
            csv_to_cert[proj_csv] = cert['nombre']
            break
        kw = False
        if 'FORMENTERA' in proj_upper and 'FORMENTERA' in cert_upper:
            if ('12' in proj_upper and '12' in cert_upper) or ('14' in proj_upper and '14' in cert_upper):
                if ('JOAQUIN' in proj_upper and 'JOAQUIN' in cert_upper) or ('JUANMA' in proj_upper and 'JUANMA' in cert_upper):
                    kw = True
        # LA PARTE DIFICIL is now split into ALICANTE + MURCIA
        if 'EDIFICIO ALICANTE' in proj_upper and 'ALICANTE' in cert_upper and 'PARTE DIFICIL' in cert_upper: kw = True
        if 'CUARTEL DE ARTILLERIA' in proj_upper and 'MURCIA' in cert_upper and 'PARTE DIFICIL' in cert_upper: kw = True
        # Additional: match by EDIFICIO keyword in cert name
        if 'EDIFICIO' in proj_upper and 'EDIFICIO' in cert_upper: kw = True
        if 'CUARTEL' in proj_upper and 'CUARTEL' in cert_upper: kw = True
        if 'SANTA ROSA' in proj_upper and 'SANTA ROSA' in cert_upper: kw = True
        if 'CARTAGENA' in proj_upper and 'CARTAGENA' in cert_upper: kw = True
        if 'CASTILLO' in proj_upper and 'CASTILLO' in cert_upper: kw = True
        if 'GARAJE' in proj_upper and 'GARAJE' in cert_upper: kw = True
        if 'PEREAMAR' in proj_upper and 'PEREAMAR' in cert_upper: kw = True
        if 'CBS' in proj_upper and 'CBS' in cert_upper: kw = True
        if 'CEMENTERIO' in proj_upper and 'CEMENTERIO' in cert_upper: kw = True
        if 'BARINAS' in proj_upper and 'BARINAS' in cert_upper: kw = True
        if 'ALBERCA' in proj_upper and 'ALBERCA' in cert_upper: kw = True
        if 'PLAZA CIRCULAR' in proj_upper and 'PLAZA CIRCULAR' in cert_upper: kw = True
        if 'HELENA' in proj_upper and 'HELENA' in cert_upper: kw = True
        if 'A-13' in proj_upper and 'A-13' in cert_upper: kw = True
        if 'CARLA' in proj_upper and 'CARLA' in cert_upper: kw = True
        if 'PADRE TRINI' in proj_upper and 'PADRE TRINI' in cert_upper: kw = True
        if 'PISO IBI' in proj_upper and 'PISO IBI' in cert_upper: kw = True
        if 'LUC2' in proj_upper and 'LUC2' in cert_upper: kw = True
        if 'MEJORA DEL VALLADO' in proj_upper and 'CASTILLO' in cert_upper: kw = True
        if ('ANGEL' in proj_upper or 'NGEL' in proj_upper) and 'HELENA' in cert_upper and 'CARMEN' in proj_upper: kw = True
        if kw:
            csv_to_cert[proj_csv] = cert['nombre']
            break

print("\n=== MAPEO ===")
for csv_proj, cert_proj in csv_to_cert.items():
    d = directos_por_proyecto[csv_proj]
    print("  %s -> %s (%s EUR)" % (csv_proj[:50], cert_proj[:30], "{:>10,.2f}".format(d['total'])))

sin_cert_total = 0
print("\nSin mapeo:")
for proj in directos_por_proyecto:
    if proj not in csv_to_cert:
        d = directos_por_proyecto[proj]
        sin_cert_total += d['total']
        print("  %s (%s EUR)" % (proj[:50], "{:>10,.2f}".format(d['total'])))
print("  TOTAL sin cert: %s EUR" % "{:,.2f}".format(sin_cert_total))

# ===== 5. CONSOLIDAR - ALL PROJECTS (with and without certification) =====

# Build reverse mapping: cert_name -> list of csv_proj names
cert_to_csvs = {}
for csv_proj, cert_mapped in csv_to_cert.items():
    if cert_mapped not in cert_to_csvs:
        cert_to_csvs[cert_mapped] = []
    cert_to_csvs[cert_mapped].append(csv_proj)

# Collect ALL unique project names: certified + CSV without cert
all_project_names = []
# First: certified projects
for cert in certificaciones:
    all_project_names.append(cert['nombre'])
# Then: CSV projects without certification
for csv_proj in directos_por_proyecto:
    if csv_proj not in csv_to_cert:
        all_project_names.append(csv_proj)

# Build a lookup of cert data
cert_lookup = {c['nombre']: c for c in certificaciones}

# Match MO entries directly to project display names
# For certified projects: use cert name
# For non-certified: use CSV project name
# For ALBERTO Y EVA: use its own name

# Build a list of all display names
all_display_names = [p['nombre'] for p in certificaciones]  # cert names first
for csv_proj in directos_por_proyecto:
    if csv_proj not in csv_to_cert:
        all_display_names.append(csv_proj)
# Add ALBERTO Y EVA if present in MO
has_alberto = any('ALBERTO' in m['proyecto'].upper() and 'EVA' in m['proyecto'].upper() for m in mano_obra_all)
if has_alberto:
    all_display_names.append('ALBERTO Y EVA')
    all_project_names.append('ALBERTO Y EVA')  # Also add to project list for consolidation

def match_mo_to_project(mo_upper, display_names):
    """Match a MO entry to a display project name."""
    for dname in display_names:
        dname_upper = dname.upper()
        # Direct substring match
        if dname_upper in mo_upper or mo_upper in dname_upper:
            return dname
        # Special keyword matches
        if 'FORMENTERA' in mo_upper and 'FORMENTERA' in dname_upper:
            if ('12' in mo_upper and '12' in dname_upper) or ('14' in mo_upper and '14' in dname_upper):
                return dname
            if ('JOAQUIN' in mo_upper and 'JOAQUIN' in dname_upper) or ('JUANMA' in mo_upper and 'JUANMA' in dname_upper):
                return dname
        if 'CARTAGENA' in mo_upper and 'CARTAGENA' in dname_upper:
            return dname
        if 'PARTE DIFICIL' in mo_upper and 'PARTE DIFICIL' in dname_upper:
            # Distinguish ALICANTE vs MURCIA when split
            if 'ALICANTE' in mo_upper and 'ALICANTE' in dname_upper:
                return dname
            if 'MURCIA' in mo_upper and 'MURCIA' in dname_upper:
                return dname
            # Generic fallback if only one PARTE DIFICIL exists
            if 'ALICANTE' not in dname_upper and 'MURCIA' not in dname_upper:
                return dname
            # If we have both ALICANTE and MURCIA certs but MO only says PARTE DIFICIL
            # Try to match by context (if MO says ALICANTE explicitly)
            continue  # Try next display name
        if 'CEMENTERIO' in mo_upper and 'CEMENTERIO' in dname_upper:
            return dname
        if 'PEREAMAR' in mo_upper and 'PEREAMAR' in dname_upper:
            return dname
        if ('ANGEL' in mo_upper or 'NGEL' in mo_upper) and 'HELENA' in dname_upper and 'CARMEN' in mo_upper:
            return dname
        if 'MURO VECINO' in mo_upper and 'MURO VECINO' in dname_upper:
            return dname
        if 'OBRA CAMPO' in mo_upper and 'CAMPO' in dname_upper and 'ARANTXA' in dname_upper:
            return dname
        if 'ALBERTO' in mo_upper and 'EVA' in mo_upper and 'ALBERTO' in dname_upper:
            return dname
    return None

mo_to_project = {}
for mo in mano_obra_all:
    matched_name = match_mo_to_project(mo['proyecto'].upper(), all_display_names)
    mo_to_project[id(mo)] = matched_name

print("\n=== MO -> Project mapping ===")
for mo in mano_obra_all:
    proj_match = mo_to_project.get(id(mo), None)
    print("  %-45s -> %s" % (mo['proyecto'][:45], proj_match[:35] if proj_match else 'NINGUNO'))

# Build consolidated data for ALL projects
proyectos_data = []
for pname in all_project_names:
    cert_data = cert_lookup.get(pname, None)
    cert_importe = cert_data['importe'] if cert_data else 0.0
    pct = cert_importe / total_cert if total_cert > 0 and cert_data else 0.0
    
    # Direct expenses: sum all CSV projects mapped to this cert (or this CSV project directly)
    direct_total = 0.0
    direct_count = 0
    if cert_data:
        # Certified project: sum all CSV projects mapped to it
        for csv_proj in cert_to_csvs.get(pname, []):
            direct_total += directos_por_proyecto[csv_proj]['total']
            direct_count += directos_por_proyecto[csv_proj]['count']
    elif pname in directos_por_proyecto:
        # Non-certified project: use its own direct expenses
        direct_total = directos_por_proyecto[pname]['total']
        direct_count = directos_por_proyecto[pname]['count']
    
    # Prorrateo: only for certified projects
    prorrateo = pct * total_gastos_comunes
    
    # Mano de obra: sum all MO entries that map to this project
    mo_horas = 0
    mo_coste = 0.0
    for mo in mano_obra_all:
        proj_match = mo_to_project.get(id(mo), None)
        if proj_match == pname:
            mo_horas += mo['horas']
            mo_coste += mo_cost(mo)
    
    total_coste = direct_total + prorrateo + mo_coste
    margen = cert_importe - total_coste
    margen_pct = (margen / cert_importe * 100) if cert_importe > 0 else 0
    
    proyectos_data.append({
        'nombre': pname,
        'certificacion': cert_importe,
        'pct': pct,
        'gastos_directos': direct_total,
        'direct_count': direct_count,
        'prorrateo': prorrateo,
        'mano_obra_horas': mo_horas,
        'mano_obra_coste': mo_coste,
        'total_coste': total_coste,
        'margen': margen,
        'margen_pct': margen_pct,
        'has_cert': cert_data is not None,
    })

# Sort: certified first (by cert desc), then non-certified (by direct expenses desc)
proyectos_data.sort(key=lambda x: (0 if x['has_cert'] else 1, -x['certificacion'] if x['has_cert'] else -x['gastos_directos']))

sum_cert = sum(p['certificacion'] for p in proyectos_data)
sum_directos = sum(p['gastos_directos'] for p in proyectos_data)
sum_prorrateo = sum(p['prorrateo'] for p in proyectos_data)
sum_mo = sum(p['mano_obra_coste'] for p in proyectos_data)
sum_horas = sum(p['mano_obra_horas'] for p in proyectos_data)
sum_coste = sum(p['total_coste'] for p in proyectos_data)
sum_margen = sum(p['margen'] for p in proyectos_data)
sin_cert_proyectos = [p for p in proyectos_data if not p['has_cert']]
sum_sin_cert = sum(p['gastos_directos'] + p['mano_obra_coste'] for p in sin_cert_proyectos)

print("\n" + "="*80)
print("RESUMEN FINAL:")
print("  Certificaciones:    %s EUR (%d proyectos)" % ("{:>12,.2f}".format(sum_cert), len(proyectos_data)))
print("  Gastos Directos:    %s EUR" % "{:>12,.2f}".format(sum_directos))
print("  Prorrateo:          %s EUR" % "{:>12,.2f}".format(sum_prorrateo))
print("  Mano de Obra:       %s EUR (%d horas)" % ("{:>12,.2f}".format(sum_mo), sum_horas))
print("  TOTAL COSTE:        %s EUR" % "{:>12,.2f}".format(sum_coste))
if sum_cert > 0:
    print("  MARGEN:             %s EUR (%.1f%%)" % ("{:>12,.2f}".format(sum_margen), sum_margen/sum_cert*100))
print("  Sin certificacion:  %s EUR" % "{:>12,.2f}".format(sum_sin_cert))

# ===== 6. GENERAR HTML =====
def fmt(val):
    if val < 0:
        return "-" + fmt(-val)
    s = "{:,.2f}".format(val)
    s = s.replace(',', 'X').replace('.', ',').replace('X', '.')
    return s

def fmt_short(val):
    if val < 0:
        return "-" + fmt_short(-val)
    s = "{:,.0f}".format(val)
    s = s.replace(',', 'X').replace('.', ',').replace('X', '.')
    return s

# Build margen table rows - ALL projects in one unified table
margen_rows = ""
for p in proyectos_data:
    margen_class = "pos" if p['margen'] >= 0 else "neg"
    bar_width = min(abs(p['margen_pct']), 100)
    bar_color = "#2ecc71" if p['margen'] >= 0 else "#e74c3c"
    # Style: non-certified projects get a subtle different background
    row_style = '' if p['has_cert'] else ' style="background:#fff8f0"'
    cert_label = fmt(p['certificacion']) if p['has_cert'] else '<span style="color:#999">-</span>'
    prorrateo_label = fmt(p['prorrateo']) if p['has_cert'] else '<span style="color:#999">-</span>'
    pct_label = '%.1f%%' % p['margen_pct'] if p['has_cert'] else '-'
    name_extra = '' if p['has_cert'] else ' <span style="font-size:0.7rem;color:#999">(s/c)</span>'
    margen_rows += '<tr%s>\n' % row_style
    safe_name = p['nombre'].replace('"', '&quot;')
    # Shorten display name for table (keep full in tooltip)
    _dn = p['nombre']
    if len(_dn) > 28:
        _dn = _dn[:26].rstrip() + '..'
    margen_rows += '  <td title="%s"><strong>%s</strong>%s <button onclick="event.stopPropagation();generateObraPDF(\'%s\')" title="PDF de esta obra" style="background:#D4742C;color:white;border:none;border-radius:3px;padding:2px 6px;cursor:pointer;font-size:0.65rem;font-weight:700;margin-left:4px;vertical-align:middle">PDF</button></td>\n' % (p['nombre'], _dn, name_extra, safe_name)
    margen_rows += '  <td class="num">%s</td>\n' % cert_label
    margen_rows += '  <td class="num">%s</td>\n' % fmt(p['gastos_directos'])
    margen_rows += '  <td class="num">%s</td>\n' % prorrateo_label
    margen_rows += '  <td class="num">%s</td>\n' % fmt(p['mano_obra_coste'])
    margen_rows += '  <td class="num"><strong>%s</strong></td>\n' % fmt(p['total_coste'])
    margen_rows += '  <td class="num">%s</td>\n' % pct_label
    margen_rows += '  <td><span class="%s" style="font-weight:700">%s EUR</span>\n' % (margen_class, fmt(p['margen']))
    margen_rows += '      <div class="progress-bar" style="margin-top:4px"><div class="progress-fill" style="width:%.0f%%;background:%s"></div></div>\n' % (bar_width, bar_color)
    margen_rows += '  </td>\n</tr>\n'

# Total row
margen_rows += '<tr class="total-row">\n'
margen_rows += '  <td>TOTAL GENERAL (%d proyectos)</td>\n' % len(proyectos_data)
margen_rows += '  <td class="num">%s</td>\n' % fmt(sum_cert)
margen_rows += '  <td class="num">%s</td>\n' % fmt(sum_directos)
margen_rows += '  <td class="num">%s</td>\n' % fmt(sum_prorrateo)
margen_rows += '  <td class="num">%s</td>\n' % fmt(sum_mo)
margen_rows += '  <td class="num"><strong>%s</strong></td>\n' % fmt(sum_coste)
margen_rows += '  <td class="num"></td>\n'
if sum_cert > 0:
    margen_rows += '  <td><span class="pos" style="font-weight:700">%s EUR (%.1f%%)</span></td>\n' % (fmt(sum_margen), sum_margen/sum_cert*100)
else:
    margen_rows += '  <td><span class="neg" style="font-weight:700">%s EUR</span></td>\n' % fmt(sum_margen)
margen_rows += '</tr>\n'

# Prorrateo table
prorr_rows = ""
for p in proyectos_data:
    prorr_rows += '<tr><td>%s</td><td class="num">%s</td><td class="num">%.2f%%</td><td class="num">%s</td></tr>\n' % (p['nombre'], fmt(p['certificacion']), p['pct']*100, fmt(p['prorrateo']))
prorr_rows += '<tr class="total-row"><td>TOTAL</td><td class="num">%s</td><td class="num">100.00%%</td><td class="num">%s</td></tr>\n' % (fmt(sum_cert), fmt(total_gastos_comunes))

# Chart data
chart_labels = json.dumps([p['nombre'][:25] for p in proyectos_data], ensure_ascii=False)
chart_cert = json.dumps([round(p['certificacion'], 2) for p in proyectos_data])
chart_directos = json.dumps([round(p['gastos_directos'], 2) for p in proyectos_data])
chart_prorr = json.dumps([round(p['prorrateo'], 2) for p in proyectos_data])
chart_mo_vals = json.dumps([round(p['mano_obra_coste'], 2) for p in proyectos_data])
chart_margen = json.dumps([round(p['margen'], 2) for p in proyectos_data])
chart_margen_colors = json.dumps(["#2ecc71" if p['margen'] >= 0 else "#e74c3c" for p in proyectos_data])

# Mano de obra table (combined 2025+2026 like original)
mo_rows = ""
mo_2026 = [m for m in mano_obra_all if m['year'] == 2026]
mo_2025 = [m for m in mano_obra_all if m['year'] == 2025]

for m in mo_2026:
    tarifa_str = '%d,00' % m['tarifa'] if m['tarifa'] > 0 else 'N/A'
    mo_rows += '<tr><td>%s</td><td class="num">%d</td><td class="num">%s</td><td class="num">%s</td></tr>\n' % (m['proyecto'], m['horas'], tarifa_str, fmt(mo_cost(m)))

h26 = sum(m['horas'] for m in mo_2026)
c26 = sum(mo_cost(m) for m in mo_2026)
mo_rows += '<tr class="total-row"><td>SUBTOTAL 2026</td><td class="num">%d</td><td class="num"></td><td class="num">%s</td></tr>\n' % (h26, fmt(c26))

if mo_2025:
    mo_rows += '<tr><td colspan="4" style="background:#f0f0f0;font-weight:600;padding-top:14px">ANYO 2025</td></tr>\n'
    for m in mo_2025:
        tarifa_str = '%d,00' % m['tarifa'] if m['tarifa'] > 0 else 'N/A'
        mo_rows += '<tr><td>%s</td><td class="num">%d</td><td class="num">%s</td><td class="num">%s</td></tr>\n' % (m['proyecto'], m['horas'], tarifa_str, fmt(mo_cost(m)))
    h25 = sum(m['horas'] for m in mo_2025)
    c25 = sum(mo_cost(m) for m in mo_2025)
    mo_rows += '<tr class="total-row"><td>SUBTOTAL 2025</td><td class="num">%d</td><td class="num"></td><td class="num">%s</td></tr>\n' % (h25, fmt(c25))
    mo_rows += '<tr class="total-row" style="background:#E8E4DC;border-top:3px solid #2B4C6F"><td>TOTAL COMBINADO</td><td class="num">%d</td><td class="num"></td><td class="num">%s</td></tr>\n' % (h26+h25, fmt(c26+c25))

# GG detail
with open(csv_path, 'r', encoding='latin-1') as f:
    content = f.read()
reader = csv.DictReader(io.StringIO(content), delimiter=';')
gg_rows_html = ""
veh_rows_html = ""
for row in reader:
    proyecto = ''
    importe_str = '0'
    titulo = ''
    proveedor = ''
    fecha = ''
    codigo = ''
    for key, val in row.items():
        if key is None:
            continue
        kl = key.lower()
        val = val.strip() if val else ''
        if 'proyecto' == kl: proyecto = val.upper()
        elif 'importe' == kl: importe_str = val
        elif 't' in kl and 'tulo' in kl: titulo = val
        elif 'proveedor' in kl and 'id' not in kl: proveedor = val
        elif 'fecha' in kl and 'imputa' in kl: fecha = val
        elif 'digo' in kl or 'odigo' in kl: codigo = val
    
    importe = parse_euro_amount(importe_str)
    imp_class = "num neg" if importe < 0 else "num "
    
    if 'GASTOS GENERALES' in proyecto:
        cat = '<span style="background:#3498db;color:white;padding:2px 8px;border-radius:10px;font-size:0.75rem">G.Generales</span>'
        gg_rows_html += '<tr><td style="font-size:0.78rem">%s</td><td>%s</td><td>%s</td><td>%s</td><td class="%s">%s</td><td style="font-size:0.8rem">%s</td></tr>\n' % (codigo, fecha, titulo, cat, imp_class, fmt(importe), proveedor[:40])
    elif 'VEHICULOS' in proyecto or 'VEH\u00cdCULOS' in proyecto:
        cat = '<span style="background:#e67e22;color:white;padding:2px 8px;border-radius:10px;font-size:0.75rem">Vehiculo</span>'
        veh_rows_html += '<tr><td style="font-size:0.78rem">%s</td><td>%s</td><td>%s</td><td>%s</td><td class="%s">%s</td><td style="font-size:0.8rem">%s</td></tr>\n' % (codigo, fecha, titulo, cat, imp_class, fmt(importe), proveedor[:40])

# Build 'Facturas por Obra' table rows
facturas_obra_rows_html = ''
for f in all_facturas_dir:
    imp_class = 'num neg' if f['importe'] < 0 else 'num'
    facturas_obra_rows_html += '<tr><td style="font-size:0.78rem">%s</td><td style="font-size:0.78rem">%s</td><td>%s</td><td style="font-size:0.8rem">%s</td><td style="font-size:0.8rem">%s</td><td>%s</td><td class="%s">%s</td></tr>\n' % (
        f['proyecto'][:50], f['codigo'], f['fecha'],
        f['titulo'][:35], f['proveedor'][:35], f['estado'],
        imp_class, fmt(f['importe'])
    )

# Calculate top 10 suppliers by total amount
supplier_totals = collections.defaultdict(float)
for f in all_facturas_dir:
    prov = f['proveedor'].strip() if f['proveedor'] else '(Sin proveedor)'
    supplier_totals[prov] += f['importe']
top10_suppliers = sorted(supplier_totals.items(), key=lambda x: -abs(x[1]))[:10]
top10_labels = [s[0][:30] for s in top10_suppliers]
top10_values = [round(abs(s[1]), 2) for s in top10_suppliers]

# Top 10 projects by total amount (direct expenses)
top10_proj = sorted(directos_por_proyecto.items(), key=lambda x: -abs(x[1]['total']))[:10]
top10_proj_labels = [p[0][:30] for p in top10_proj]
top10_proj_values = [round(abs(p[1]['total']), 2) for p in top10_proj]

# Full names for filtering
top10_sup_full = [s[0] for s in top10_suppliers]
top10_proj_full = [p[0] for p in top10_proj]

# GG categories
from collections import Counter
gg_cats = Counter()
for row_str in gg_rows_html.split('\n'):
    pass  # Already computed from gg_rows list

# Recalculate GG categories from detail
gg_cat_map = {
    'Combustible': 0, 'Alquiler': 0, 'Parking': 0, 'Ropa Laboral': 0,
    'Asesoria': 0, 'Telecomunicaciones': 0, 'Preencion': 0, 'Material': 0,
    'Lavado': 0, 'Seguros': 0, 'Otros': 0
}

# Re-read for categories
with open(csv_path, 'r', encoding='latin-1') as f:
    content = f.read()
reader = csv.DictReader(io.StringIO(content), delimiter=';')
for row in reader:
    proyecto = ''
    importe_str = '0'
    titulo = ''
    for key, val in row.items():
        if key is None:
            continue
        kl = key.lower()
        val = val.strip() if val else ''
        if 'proyecto' == kl: proyecto = val.upper()
        elif 'importe' == kl: importe_str = val
        elif 't' in kl and 'tulo' in kl: titulo = val.upper()
    
    if 'GASTOS GENERALES' not in proyecto:
        continue
    importe = parse_euro_amount(importe_str)
    
    if any(w in titulo for w in ['REPOSTAJE', 'PLENOIL', 'GASO', 'COMBUSTIBLE', 'DIESEL', 'GASOLINA']):
        gg_cat_map['Combustible'] += importe
    elif 'ALQUILER' in titulo:
        gg_cat_map['Alquiler'] += importe
    elif 'PARKING' in titulo:
        gg_cat_map['Parking'] += importe
    elif any(w in titulo for w in ['ROPA', 'CALZADO', 'BOTAS']):
        gg_cat_map['Ropa Laboral'] += importe
    elif any(w in titulo for w in ['ASESOR', 'CONTAB', 'FISCAL']):
        gg_cat_map['Asesoria'] += importe
    elif any(w in titulo for w in ['TEL', 'DIGI', 'ORANGE', 'MOVISTAR', 'TELEFON']):
        gg_cat_map['Telecomunicaciones'] += importe
    elif any(w in titulo for w in ['RECONOCIMIENTO', 'MEDICO', 'PREVENCION', 'MEDICINA']):
        gg_cat_map['Preencion'] += importe
    elif any(w in titulo for w in ['MATERIAL', 'HERRAM', 'COMPR', 'BROCA', 'CORTAD', 'DISCO', 'MALLA', 'PINTURA', 'TORNILL', 'LLAVE']):
        gg_cat_map['Material'] += importe
    elif 'LAVADO' in titulo:
        gg_cat_map['Lavado'] += importe
    elif any(w in titulo for w in ['SEGURO', 'SEGUROMERC']):
        gg_cat_map['Seguros'] += importe
    else:
        gg_cat_map['Otros'] += importe

# Remove zero categories
gg_cat_labels = json.dumps([k for k, v in gg_cat_map.items() if v > 0])
gg_cat_values = json.dumps([round(v, 2) for v in gg_cat_map.values() if v > 0])

# Veh categories
veh_cat_map = Counter()
with open(csv_path, 'r', encoding='latin-1') as f:
    content = f.read()
reader = csv.DictReader(io.StringIO(content), delimiter=';')
for row in reader:
    proyecto = ''
    importe_str = '0'
    titulo = ''
    for key, val in row.items():
        if key is None:
            continue
        kl = key.lower()
        val = val.strip() if val else ''
        if 'proyecto' == kl: proyecto = val.upper()
        elif 'importe' == kl: importe_str = val
        elif 't' in kl and 'tulo' in kl: titulo = val.upper()
    
    if 'VEHICULOS' not in proyecto and 'VEH\u00cdCULOS' not in proyecto:
        continue
    importe = parse_euro_amount(importe_str)
    if any(w in titulo for w in ['REPOSTAJE', 'GASO', 'COMBUSTIBLE', 'DIESEL']):
        veh_cat_map['Combustible'] += importe
    elif 'ITV' in titulo:
        veh_cat_map['ITV'] += importe
    elif 'ASEGUR' in titulo:
        veh_cat_map['Seguro'] += importe
    elif any(w in titulo for w in ['REPARA', 'MANTEN']):
        veh_cat_map['Reparacion'] += importe
    elif 'LAVADO' in titulo:
        veh_cat_map['Lavado'] += importe
    elif 'NEUMATICOS' in titulo or 'NEUM' in titulo:
        veh_cat_map['Neumaticos'] += importe
    else:
        veh_cat_map['Otros'] += importe

veh_cat_labels = json.dumps(list(veh_cat_map.keys()))
veh_cat_values = json.dumps([round(v, 2) for v in veh_cat_map.values()])

margen_color = "#27ae60" if sum_margen >= 0 else "#e74c3c"

# Build the HTML
lines = []
lines.append('<!DOCTYPE html>')
lines.append('<html lang="es">')
lines.append('<head>')
lines.append('<meta charset="UTF-8">')
lines.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
lines.append('<title>ECO STRUCT - Dashboard Financiero</title>')
lines.append('<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.4/dist/chart.umd.min.js"></script>')
lines.append('<script src="https://cdnjs.cloudflare.com/ajax/libs/jspdf/2.5.1/jspdf.umd.min.js"></script>')
lines.append('<style>')
lines.append(':root { --primary:#2B3A4E; --accent:#D4742C; --highlight:#e94560; --success:#2ecc71; --warning:#f39c12; --info:#5BA3C9; --navy:#2B4C6F; --bg-page:#F5F0E8; --bg-card:#FFFFFF; }')
lines.append('* { margin:0; padding:0; box-sizing:border-box; }')
lines.append("body { font-family:'Inter','Segoe UI',Tahoma,sans-serif; background:var(--bg-page); color:#2B3A4E; }")
lines.append(".header { background:white; color:#2B3A4E; padding:20px 40px; display:flex; justify-content:space-between; align-items:center; border-bottom:3px solid var(--accent); }")
lines.append(".header h1 { font-size:1.6rem; font-weight:700; color:#2B3A4E; letter-spacing:-0.5px; }")
lines.append(".header .date { font-size:0.85rem; color:#5a6a7a; }")
lines.append(".container { max-width:1500px; margin:0 auto; padding:24px; }")
lines.append(".kpi-row { display:grid; grid-template-columns:repeat(auto-fit,minmax(200px,1fr)); gap:14px; margin-bottom:24px; }")
lines.append(".kpi-card { background:white; border-radius:8px; padding:18px 22px; box-shadow:0 2px 8px rgba(0,0,0,0.06); transition:transform .2s; border-left:5px solid var(--accent); }")
lines.append(".kpi-card:hover { transform:translateY(-2px); box-shadow:0 4px 12px rgba(0,0,0,0.12); }")
lines.append(".kpi-card .label { font-size:0.72rem; text-transform:uppercase; color:#5a6a7a; letter-spacing:1.5px; font-weight:600; margin-bottom:6px; }")
lines.append(".kpi-card .value { font-size:1.6rem; font-weight:800; color:#2B3A4E; font-variant-numeric:tabular-nums; }")
lines.append(".kpi-card .sub { font-size:0.75rem; color:#8899aa; margin-top:4px; }")
lines.append(".section { background:white; border-radius:10px; padding:24px; margin-bottom:24px; box-shadow:0 2px 8px rgba(0,0,0,0.06); }")
lines.append(".section-title { font-size:0.95rem; font-weight:700; color:#2B3A4E; margin-bottom:14px; text-transform:uppercase; letter-spacing:1.5px; }")
lines.append(".grid-2 { display:grid; grid-template-columns:1fr 1fr; gap:24px; }")
lines.append(".table-wrapper { overflow-x:auto; }")
lines.append("table { width:100%; border-collapse:collapse; font-size:0.82rem; }")
lines.append("thead th { background:var(--navy); color:white; padding:10px 14px; text-align:left; font-weight:600; white-space:nowrap; font-size:0.78rem; text-transform:uppercase; letter-spacing:0.5px; }")
lines.append("thead th:first-child { border-radius:6px 0 0 0; }")
lines.append("thead th:last-child { border-radius:0 6px 0 0; }")
lines.append("tbody td { padding:9px 14px; border-bottom:1px solid #E8E4DC; white-space:nowrap; }")
lines.append("#projTable { font-size:0.68rem; width:100%; }")
lines.append("#projTable thead th { padding:5px 4px; font-size:0.62rem; letter-spacing:0.2px; }")
lines.append("#projTable tbody td { padding:4px 4px; white-space:nowrap; }")
lines.append("#projTable tbody td:first-child { white-space:normal; overflow:hidden; text-overflow:ellipsis; max-width:140px; }")
lines.append("#projTable .progress-bar { height:3px; margin-top:2px; }")
lines.append("tbody tr:hover { background:#F0EDE6; }")
lines.append("tbody tr:nth-child(even) { background:#FAF8F4; }")
lines.append(".num { text-align:right; font-variant-numeric:tabular-nums; }")
lines.append(".pos { color:#2B8A3E; font-weight:700; }")
lines.append(".neg { color:#C0392B; font-weight:700; }")
lines.append(".total-row { font-weight:700; background:#E8E4DC !important; border-top:2px solid var(--navy); }")
lines.append(".tabs { display:flex; gap:4px; margin-bottom:20px; flex-wrap:wrap; }")
lines.append(".tab-btn { padding:10px 20px; border:none; border-radius:0; cursor:pointer; font-size:0.85rem; font-weight:600; background:transparent; color:#5a6a7a; transition:all .2s; border-bottom:3px solid transparent; text-transform:uppercase; letter-spacing:0.5px; }")
lines.append(".tab-btn.active { color:var(--accent); border-bottom:3px solid var(--accent); background:transparent; }")
lines.append(".tab-btn:hover:not(.active) { color:#2B3A4E; background:#EDE9E1; }")
lines.append(".tab-content { display:none; } .tab-content.active { display:block; }")
lines.append(".search-box { padding:8px 14px; border:2px solid #D5D0C8; border-radius:6px; font-size:0.83rem; width:260px; outline:none; background:white; }")
lines.append(".search-box:focus { border-color:var(--accent); }")
lines.append(".toolbar { display:flex; justify-content:space-between; align-items:center; margin-bottom:14px; flex-wrap:wrap; gap:10px; }")
lines.append(".progress-bar { background:#E8E4DC; border-radius:4px; height:6px; overflow:hidden; width:100%; }")
lines.append(".progress-fill { height:100%; border-radius:4px; }")
lines.append(".chart-container { position:relative; height:350px; }")
lines.append("@media(max-width:900px) { .grid-2 { grid-template-columns:1fr; } }")
lines.append('</style>')
lines.append('</head>')
lines.append('<body>')

# Generate base64 logo (rotated 180 to face upward)
import base64 as _b64
import io as _io
from PIL import Image as _Img
_logo_orig = _Img.open(r'C:\NAS\03PUBLI\LOGO\LOGO 1\For Web\png\symbol.png')
_logo_rot = _logo_orig.rotate(180)
_logo_buf = _io.BytesIO()
_logo_rot.save(_logo_buf, format='PNG')
_logo_b64 = 'data:image/png;base64,' + _b64.b64encode(_logo_buf.getvalue()).decode()

# Header
lines.append('<div class="header">')
lines.append('  <div>')
lines.append('    <h1>ECO STRUCT - Dashboard Financiero</h1>')
lines.append('    <div class="date">Datos a fecha de 31/08/2026 | Gastos directos del CSV de Holded</div>')
lines.append('  </div>')
lines.append('  <div style="display:flex;align-items:center;gap:12px"><img src="' + _logo_b64 + '" alt="ECO STRUCT" style="height:50px;width:auto"><div style="text-align:right"><div style="font-size:1.25rem;font-weight:800;color:#2B3A4E;letter-spacing:0.5px">Constructive Ecosen</div><div style="font-size:1.1rem;font-weight:700;color:#5a6a7a">Spain 2.3</div></div></div>')
lines.append('</div>')

lines.append('<div class="container">')

# KPIs
lines.append('<div class="kpi-row">')
kpi_cards = [
    ('#3498db', 'Certificaciones', fmt_short(sum_cert) + ' EUR', '%d proyectos' % len(proyectos_data), 0),
    ('#e94560', 'Gastos Directos', fmt_short(total_directos_csv) + ' EUR', '%d facturas por obra (%d sin certif.)' % (total_facturas_dir, len(sin_cert_proyectos)), 5),
    ('#f39c12', 'Gastos Comunes', fmt_short(total_gastos_comunes) + ' EUR', 'GG %s + VEH %s' % (fmt_short(gg_total), fmt_short(veh_total)), 2),
    ('#9b59b6', 'Mano de Obra', fmt_short(sum_mo) + ' EUR', '%s horas' % "{:,}".format(sum_horas), 4),
]
for color, label, value, sub, tab_idx in kpi_cards:
    card_js = "document.querySelectorAll('.tab-btn')[%d].click()" % tab_idx
    lines.append('  <div class="kpi-card" style="border-left-color:%s;cursor:pointer" onclick="%s"><div class="label">%s</div><div class="value">%s</div><div class="sub">%s</div></div>' % (color, card_js, label, value, sub))
lines.append('  <div class="kpi-card" style="border-left-color:%s"><div class="label">MARGEN TOTAL</div><div class="value" style="color:%s">%s EUR</div><div class="sub">%.1f%% sobre certificacion</div></div>' % (margen_color, margen_color, fmt(sum_margen), sum_margen/sum_cert*100 if sum_cert > 0 else 0))
lines.append('</div>')

# Tabs
lines.append('<div class="tabs">')
lines.append('  <button class="tab-btn active" onclick="showTab(\'resumen\')">Resumen General</button>')
lines.append('  <button class="tab-btn" onclick="showTab(\'margen\')">Obras</button>')
lines.append('  <button class="tab-btn" onclick="showTab(\'prorrateo\')">Prorrateo Gastos</button>')
lines.append('  <button class="tab-btn" onclick="showTab(\'gastosGen\')">Gastos Generales (%d+%d)</button>' % (gg_count, veh_count))
lines.append('  <button class="tab-btn" onclick="showTab(\'manoObra\')">Mano de Obra</button>')
lines.append('  <button class="tab-btn" onclick="showTab(\'facturasObra\')">Facturas por Obra (%d)</button>' % total_facturas_dir)
# Build project options for the PDF selector
pdf_proj_options = '<option value="">Todas las obras</option>'
for p in proyectos_data:
    pdf_proj_options += '<option value="%s">%s</option>' % (p['nombre'].replace('"', '&quot;'), p['nombre'])

lines.append('  <div style="margin-left:auto;display:flex;align-items:center;gap:6px">')
lines.append('    <select id="pdfProjectSelect" style="padding:8px 10px;border:2px solid #D4742C;border-radius:6px;font-size:0.8rem;font-weight:600;background:white;color:#2B3A4E;cursor:pointer;max-width:260px">')
lines.append('      %s' % pdf_proj_options)
lines.append('    </select>')
lines.append('    <button id="btnPDF" onclick="generatePDF()" style="background:var(--accent);color:white;border:none;border-radius:6px;padding:10px 22px;cursor:pointer;font-size:0.85rem;font-weight:700;text-transform:uppercase;letter-spacing:1px;display:flex;align-items:center;gap:6px;box-shadow:0 2px 8px rgba(212,116,44,0.3);transition:all .2s"><svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/><polyline points="10 9 9 9 8 9"/></svg> Exportar PDF</button>')
lines.append('  </div>')
lines.append('</div>')

# RESUMEN tab
lines.append('<div class="tab-content active" id="tab-resumen">')

# Summary mini-cards row
cert_pct_of_total = (sum_cert / sum_cert * 100) if sum_cert > 0 else 0
efficiency = sum_cert / sum_coste if sum_coste > 0 else 0
avg_margin_pct = sum_margen / sum_cert * 100 if sum_cert > 0 else 0
projects_with_margin = len([p for p in proyectos_data if p['has_cert'] and p['margen'] > 0])
projects_with_loss = len([p for p in proyectos_data if p['has_cert'] and p['margen'] < 0])
lines.append('  <div style="display:grid;grid-template-columns:repeat(5,1fr);gap:10px;margin-bottom:16px">')
lines.append('    <div style="background:linear-gradient(135deg,#3498db,#2980b9);color:white;border-radius:10px;padding:14px 12px;text-align:center"><div style="font-size:0.7rem;opacity:0.85;margin-bottom:4px">Eficiencia Coste/Certif.</div><div style="font-size:1.5rem;font-weight:700">%.1f%%</div></div>' % (efficiency * 100))
lines.append('    <div style="background:linear-gradient(135deg,#27ae60,#229954);color:white;border-radius:10px;padding:14px 12px;text-align:center"><div style="font-size:0.7rem;opacity:0.85;margin-bottom:4px">Obras con Beneficio</div><div style="font-size:1.5rem;font-weight:700">%d / %d</div></div>' % (projects_with_margin, len([p for p in proyectos_data if p['has_cert']])))
lines.append('    <div style="background:linear-gradient(135deg,#e74c3c,#c0392b);color:white;border-radius:10px;padding:14px 12px;text-align:center"><div style="font-size:0.7rem;opacity:0.85;margin-bottom:4px">Obras con Perdida</div><div style="font-size:1.5rem;font-weight:700">%d / %d</div></div>' % (projects_with_loss, len([p for p in proyectos_data if p['has_cert']])))
lines.append('    <div style="background:linear-gradient(135deg,#f39c12,#e67e22);color:white;border-radius:10px;padding:14px 12px;text-align:center"><div style="font-size:0.7rem;opacity:0.85;margin-bottom:4px">Facturas Directas</div><div style="font-size:1.5rem;font-weight:700">%s</div></div>' % ("{:,}".format(total_facturas_dir)))
lines.append('    <div style="background:linear-gradient(135deg,#9b59b6,#8e44ad);color:white;border-radius:10px;padding:14px 12px;text-align:center"><div style="font-size:0.7rem;opacity:0.85;margin-bottom:4px">Horas Mano de Obra</div><div style="font-size:1.5rem;font-weight:700">%s h</div></div>' % ("{:,}".format(sum_horas)))
lines.append('  </div>')

# Row 1: Donut global + Cost composition stacked bars
lines.append('  <div class="grid-2">')
lines.append('    <div class="section"><div class="section-title">Composicion Global de Costes</div><div style="text-align:center"><canvas id="chartGlobalDonut" width="300" height="300"></canvas></div></div>')
lines.append('    <div class="section"><div class="section-title">Composicion de Costes por Obra <span style="font-size:0.75rem;color:#888;font-weight:400">(clic en una barra para ver detalle)</span></div><div class="chart-container"><canvas id="chartCostStack"></canvas></div></div>')
lines.append('  </div>')

# Detail panel (hidden by default)
lines.append('  <div class="section" id="resumen-detail-panel" style="display:none;border:2px solid #2B4C6F;background:#F5F0E8;border-radius:10px">')
lines.append('    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:10px">')
lines.append('      <div class="section-title" style="margin:0" id="resumen-detail-title">COMPOSICION</div>')
lines.append('      <button onclick="document.getElementById(\'resumen-detail-panel\').style.display=\'none\'" style="background:#e74c3c;color:white;border:none;border-radius:50%;width:28px;height:28px;cursor:pointer;font-size:1rem">&times;</button>')
lines.append('    </div>')
lines.append('    <div class="grid-2">')
lines.append('      <div style="text-align:center"><canvas id="resumenDonut" width="320" height="320"></canvas></div>')
lines.append('      <div>')
lines.append('        <div id="resumen-detail-legend" style="margin-bottom:16px"></div>')
lines.append('        <div style="background:white;border:2px solid #2B4C6F;border-radius:8px;padding:16px 20px">')
lines.append('          <div id="resumen-detail-summary" style="font-family:Inter,monospace;font-size:0.9rem;line-height:1.8"></div>')
lines.append('        </div>')
lines.append('      </div>')
lines.append('    </div>')
lines.append('  </div>')

# Row 2: Margen waterfall + Certification bar chart
lines.append('  <div class="grid-2">')
lines.append('    <div class="section"><div class="section-title">Margen por Proyecto <span style="font-size:0.75rem;color:#888;font-weight:400">(clic en una barra para ver detalle)</span></div><div class="chart-container" style="height:420px"><canvas id="chartMargenBar"></canvas></div></div>')
lines.append('    <div class="section"><div class="section-title">Certificacion vs Coste Total</div><div class="chart-container" style="height:420px"><canvas id="chartComp"></canvas></div></div>')
lines.append('  </div>')

lines.append('</div>')

# Build project data JSON for click detail
projects_json_data = []
for p in proyectos_data:
    projects_json_data.append({
        'nombre': p['nombre'],
        'certificacion': round(p['certificacion'], 2),
        'gastos_directos': round(p['gastos_directos'], 2),
        'prorrateo': round(p['prorrateo'], 2),
        'mano_obra': round(p['mano_obra_coste'], 2),
        'total_coste': round(p['total_coste'], 2),
        'margen': round(p['margen'], 2),
        'margen_pct': round(p['margen_pct'], 1),
        'has_cert': p['has_cert'],
    })
projects_json = json.dumps(projects_json_data, ensure_ascii=False)

# OBRAS tab
lines.append('<div class="tab-content" id="tab-margen">')
lines.append('  <div class="section" id="obra-detail-panel" style="display:none;border:2px solid #2B4C6F;background:#F5F0E8;border-radius:10px">')
lines.append('    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:10px">')
lines.append('      <div class="section-title" style="margin:0" id="detail-title" data-nombre="">COMPOSICION</div>')
lines.append('      <div style="display:flex;gap:6px;align-items:center">')
lines.append('        <button onclick="pdfFromDetail()" style="background:#D4742C;color:white;border:none;border-radius:4px;padding:4px 10px;cursor:pointer;font-size:0.7rem;font-weight:700">PDF Obra</button>')
lines.append('        <button onclick="closeDetail()" style="background:#e74c3c;color:white;border:none;border-radius:50%;width:28px;height:28px;cursor:pointer;font-size:1rem">&times;</button>')
lines.append('      </div>')
lines.append('    </div>')
lines.append('    <div class="grid-2">')
lines.append('      <div style="text-align:center"><canvas id="detailDonut" width="320" height="320"></canvas></div>')
lines.append('      <div>')
lines.append('        <div id="detail-legend" style="margin-bottom:16px"></div>')
lines.append('        <div style="background:white;border:2px solid #2B4C6F;border-radius:8px;padding:16px 20px">')
lines.append('          <div id="detail-summary" style="font-family:Inter,monospace;font-size:0.9rem;line-height:1.8"></div>')
lines.append('        </div>')
lines.append('        <p style="font-size:0.8rem;color:#888;margin-top:12px">Haz clic en una fila de la tabla &quot;Obras&quot; para ver su desglose.</p>')
lines.append('      </div>')
lines.append('    </div>')
lines.append('  </div>')
lines.append('  <div class="section">')
lines.append('    <div class="toolbar"><div class="section-title" style="margin:0">Detalle por Proyecto</div><input type="text" class="search-box" placeholder="Buscar proyecto..." oninput="filterTable(\'projTable\',this.value)"></div>')
lines.append('    <div class="table-wrapper" style="overflow-x:hidden"><table id="projTable"><thead><tr><th>Proyecto</th><th class="num">Certif.</th><th class="num">G.Dir.</th><th class="num">Prorr.</th><th class="num">M.Obra</th><th class="num">TOTAL</th><th class="num">%C/C</th><th>Margen</th></tr></thead><tbody>')
lines.append(margen_rows)
lines.append('    </tbody></table></div>')
lines.append('  </div>')
lines.append('</div>')

# PRORRATEO tab
lines.append('<div class="tab-content" id="tab-prorrateo">')
lines.append('  <div class="section">')
lines.append('    <div class="section-title">Prorrateo de Gastos Comunes</div>')
lines.append('    <p style="font-size:0.85rem;color:#666;margin-bottom:14px">Gastos Generales (<strong>%s EUR</strong>) + Vehiculos (<strong>%s EUR</strong>) = <strong>%s EUR</strong> prorrateados segun %% de certificacion.</p>' % (fmt(gg_total), fmt(veh_total), fmt(total_gastos_comunes)))
lines.append('    <div class="grid-2"><div><div class="chart-container"><canvas id="chartProrrBar"></canvas></div></div><div><div class="chart-container"><canvas id="chartProrrPie"></canvas></div></div></div>')
lines.append('  </div>')
lines.append('  <div class="section">')
lines.append('    <div class="toolbar"><div class="section-title" style="margin:0">Tabla de Prorrateo</div><input type="text" class="search-box" placeholder="Buscar..." oninput="filterTable(\'prorrTable\',this.value)"></div>')
lines.append('    <div class="table-wrapper"><table id="prorrTable"><thead><tr><th>Proyecto</th><th class="num">Certificacion (EUR)</th><th class="num">%% del Total</th><th class="num">Gasto Prorrateado (EUR)</th></tr></thead><tbody>')
lines.append(prorr_rows)
lines.append('    </tbody></table></div>')
lines.append('  </div>')
lines.append('</div>')

# GASTOS GENERALES tab
lines.append('<div class="tab-content" id="tab-gastosGen">')
lines.append('  <div class="grid-2">')
lines.append('    <div class="section"><div class="section-title">Gastos Generales por Categoria</div><div class="chart-container"><canvas id="chartGG"></canvas></div></div>')
lines.append('    <div class="section"><div class="section-title">Vehiculos</div><div class="chart-container"><canvas id="chartVeh"></canvas></div></div>')
lines.append('  </div>')
lines.append('  <div class="section">')
lines.append('    <div class="section-title">Detalle Gastos Generales (%d facturas = %s EUR)</div>' % (gg_count, fmt(gg_total)))
lines.append('    <div class="table-wrapper" style="max-height:500px;overflow-y:auto"><table><thead><tr><th>Codigo</th><th>Fecha</th><th>Titulo</th><th>Cat.</th><th class="num">Importe</th><th>Proveedor</th></tr></thead><tbody>')
lines.append(gg_rows_html)
lines.append('    </tbody></table></div>')
lines.append('  </div>')
lines.append('  <div class="section">')
lines.append('    <div class="section-title">Detalle Vehiculos (%d facturas = %s EUR)</div>' % (veh_count, fmt(veh_total)))
lines.append('    <div class="table-wrapper" style="max-height:400px;overflow-y:auto"><table><thead><tr><th>Codigo</th><th>Fecha</th><th>Titulo</th><th>Cat.</th><th class="num">Importe</th><th>Proveedor</th></tr></thead><tbody>')
lines.append(veh_rows_html)
lines.append('    </tbody></table></div>')
lines.append('  </div>')
lines.append('</div>')

# MANO DE OBRA tab
lines.append('<div class="tab-content" id="tab-manoObra">')
lines.append('  <div class="section">')
lines.append('    <div class="section-title">Mano de Obra - Desglose por Proyecto</div>')
lines.append('    <p style="font-size:0.85rem;color:#666;margin-bottom:14px">Tarifa: 20 EUR/hora. Datos de "GASTOS MANO DE OBRA.docx" (2025 + 2026).</p>')
lines.append('    <div class="table-wrapper"><table><thead><tr><th>Proyecto</th><th class="num">Horas</th><th class="num">Tarifa (EUR/h)</th><th class="num">Coste Total (EUR)</th></tr></thead><tbody>')
lines.append(mo_rows)
lines.append('    </tbody></table></div>')
lines.append('  </div>')
lines.append('</div>')

# FACTURAS POR OBRA tab
lines.append('<div class="tab-content" id="tab-facturasObra">')
lines.append('  <div class="section">')
lines.append('    <div class="section-title">Facturas Directas por Obra (%d facturas = %s EUR)</div>' % (total_facturas_dir, fmt(total_directos_csv)))
lines.append('    <p style="font-size:0.85rem;color:#666;margin-bottom:14px">Todas las facturas del CSV excepto Gastos Generales y Vehiculos. Puedes buscar por proyecto, proveedor, concepto, etc.</p>')
lines.append('  </div>')
lines.append('  <div class="grid-2">')
lines.append('    <div class="section"><div class="section-title">Top 10 Proveedores por Importe</div><div class="chart-container" style="height:360px"><canvas id="chartTopProv"></canvas></div></div>')
lines.append('    <div class="section"><div class="section-title">Distribucion por Proyecto (Top 10)</div><div class="chart-container" style="height:360px"><canvas id="chartTopProj"></canvas></div></div>')
lines.append('  </div>')
lines.append('  <div class="section">')
lines.append('    <div class="toolbar"><div class="section-title" style="margin:0">Listado de Facturas <span id="activeFilterBadge" style="display:none;background:#e74c3c;color:white;padding:2px 10px;border-radius:12px;font-size:0.75rem;margin-left:8px"></span></div><div style="display:flex;gap:8px;align-items:center"><button id="clearFilterBtn" onclick="clearChartFilter()" style="display:none;background:#e74c3c;color:white;border:none;padding:6px 14px;border-radius:6px;cursor:pointer;font-size:0.8rem">\u2716 Limpiar filtro</button><input type="text" class="search-box" placeholder="Buscar por proyecto, proveedor, titulo, codigo..." oninput="filterTable(\'facturasObraTable\',this.value)"></div></div>')
lines.append('    <div class="table-wrapper" style="max-height:600px;overflow-y:auto"><table id="facturasObraTable"><thead><tr><th>Proyecto</th><th>Codigo</th><th>Fecha</th><th>Titulo/Concepto</th><th>Proveedor</th><th>Estado</th><th class="num">Importe</th></tr></thead><tbody>')
lines.append(facturas_obra_rows_html)
lines.append('    </tbody></table></div>')
lines.append('  </div>')
lines.append('</div>')

lines.append('</div>')  # container end

# JavaScript
lines.append('<script>')
lines.append("function showTab(id) {")
lines.append("  document.querySelectorAll('.tab-content').forEach(t => t.classList.remove('active'));")
lines.append("  document.querySelectorAll('.tab-btn').forEach(b => b.classList.remove('active'));")
lines.append("  document.getElementById('tab-'+id).classList.add('active');")
lines.append("  event.target.classList.add('active');")
lines.append("}")
lines.append("function filterTable(tableId, query) {")
lines.append("  var rows = document.getElementById(tableId).querySelectorAll('tbody tr');")
lines.append("  query = query.toLowerCase();")
lines.append("  rows.forEach(row => { row.style.display = row.textContent.toLowerCase().includes(query) ? '' : 'none'; });")
lines.append("}")
lines.append("var labels=%s;" % chart_labels)
lines.append("var certData=%s;" % chart_cert)
lines.append("var directData=%s;" % chart_directos)
lines.append("var prorrData=%s;" % chart_prorr)
lines.append("var moData=%s;" % chart_mo_vals)
lines.append("var margenData=%s;" % chart_margen)
lines.append("var margenColors=%s;" % chart_margen_colors)
lines.append("var ggCatLabels=%s;" % gg_cat_labels)
lines.append("var ggCatValues=%s;" % gg_cat_values)
lines.append("var vehCatLabels=%s;" % veh_cat_labels)
lines.append("var vehCatValues=%s;" % veh_cat_values)
lines.append("var projectData=%s;" % projects_json)
# Detail panel JS functions
js_lines = []
js_lines.append('var detailChart=null;')
js_lines.append('function showDetail(idx){')
js_lines.append('  var p=projectData[idx];')
js_lines.append("  var panel=document.getElementById('obra-detail-panel');")
js_lines.append("  panel.style.display='block';panel.scrollIntoView({behavior:'smooth'});")
js_lines.append("  document.getElementById('detail-title').setAttribute('data-nombre',p.nombre);document.getElementById('detail-title').textContent='COMPOSICION \u2014 '+p.nombre;")
js_lines.append('  var d=p.gastos_directos,pr=p.prorrateo,mo=p.mano_obra;')
js_lines.append('  if(detailChart){detailChart.destroy();}')
js_lines.append("  var ctx=document.getElementById('detailDonut').getContext('2d');")
js_lines.append("  detailChart=new Chart(ctx,{type:'doughnut',data:{labels:['Mano de obra','Facturas','Gastos generales (prorrateo)'],datasets:[{data:[mo,d,pr],backgroundColor:['#2B4C6F','#5BA3C9','#7A8B99'],borderWidth:3,borderColor:'#F5F0E8'}]},options:{responsive:false,cutout:'55%',plugins:{legend:{display:false}}}});")
js_lines.append("  var legendHtml='<div style=\"display:flex;gap:16px;flex-wrap:wrap;margin-bottom:8px\">';")
js_lines.append("  legendHtml+='<div><span style=\"display:inline-block;width:12px;height:12px;background:#8e44ad;border-radius:2px;margin-right:4px\"></span>Mano de obra: '+mo.toLocaleString('es-ES',{minimumFractionDigits:2})+' \u20ac</div>';")
js_lines.append("  legendHtml+='<div><span style=\"display:inline-block;width:12px;height:12px;background:#3498db;border-radius:2px;margin-right:4px\"></span>Facturas: '+d.toLocaleString('es-ES',{minimumFractionDigits:2})+' \u20ac</div>';")
js_lines.append("  legendHtml+='<div><span style=\"display:inline-block;width:12px;height:12px;background:#f39c12;border-radius:2px;margin-right:4px\"></span>Gastos generales (prorrateo): '+pr.toLocaleString('es-ES',{minimumFractionDigits:2})+' \u20ac</div>';")
js_lines.append("  legendHtml+='</div>';document.getElementById('detail-legend').innerHTML=legendHtml;")
js_lines.append("  var margenColor=p.margen>=0?'#27ae60':'#e74c3c';")
js_lines.append("  var certStr=p.has_cert?p.certificacion.toLocaleString('es-ES',{minimumFractionDigits:2})+' \u20ac':'Sin certificacion';")
js_lines.append("  var html2='<strong>CERTIFICADO</strong> '+certStr+' &mdash; <strong>COSTE TOTAL</strong> '+p.total_coste.toLocaleString('es-ES',{minimumFractionDigits:2})+' \u20ac';")
js_lines.append("  html2+='<br><span style=\"color:'+margenColor+';font-size:1.3rem;font-weight:700\">'+(p.margen>=0?'+':'')+p.margen.toLocaleString('es-ES',{minimumFractionDigits:2})+' \u20ac</span>';")
js_lines.append("  html2+='<br><span style=\"color:'+margenColor+'\">('+p.margen_pct+String.fromCharCode(37)+')</span>';")
js_lines.append("  document.getElementById('detail-summary').innerHTML=html2;}")
js_lines.append("function closeDetail(){document.getElementById('obra-detail-panel').style.display='none';}")
js_lines.append("document.addEventListener('DOMContentLoaded',function(){var rows=document.getElementById('projTable').querySelectorAll('tbody tr');rows.forEach(function(row,i){row.style.cursor='pointer';row.addEventListener('click',function(){showDetail(i);});row.addEventListener('mouseenter',function(){row.style.background='#e8f4fd';});row.addEventListener('mouseleave',function(){row.style.background='';});});});")
for jl in js_lines:
    lines.append(jl)
# showResumenDetail: show donut in Resumen tab when clicking a cert chart bar
lines.append("var resumenChart=null;")
lines.append("function showResumenDetail(idx){var p=projectData[idx];var panel=document.getElementById('resumen-detail-panel');panel.style.display='block';panel.scrollIntoView({behavior:'smooth'});document.getElementById('resumen-detail-title').setAttribute('data-nombre',p.nombre);document.getElementById('detail-title').textContent='COMPOSICION \u2014 '+p.nombre;var d=p.gastos_directos,pr=p.prorrateo,mo=p.mano_obra;if(resumenChart){resumenChart.destroy();}var ctx=document.getElementById('resumenDonut').getContext('2d');resumenChart=new Chart(ctx,{type:'doughnut',data:{labels:['Mano de obra','Facturas','Gastos generales (prorrateo)'],datasets:[{data:[mo,d,pr],backgroundColor:['#2B4C6F','#5BA3C9','#7A8B99'],borderWidth:3,borderColor:'#F5F0E8'}]},options:{responsive:false,cutout:'55%',plugins:{legend:{display:false}}}});var legendHtml='<div style=\"display:flex;gap:16px;flex-wrap:wrap;margin-bottom:8px\">';legendHtml+='<div><span style=\"display:inline-block;width:12px;height:12px;background:#2B4C6F;border-radius:2px;margin-right:4px\"></span>Mano de obra: '+mo.toLocaleString('es-ES',{minimumFractionDigits:2})+' \u20ac</div>';legendHtml+='<div><span style=\"display:inline-block;width:12px;height:12px;background:#5BA3C9;border-radius:2px;margin-right:4px\"></span>Facturas: '+d.toLocaleString('es-ES',{minimumFractionDigits:2})+' \u20ac</div>';legendHtml+='<div><span style=\"display:inline-block;width:12px;height:12px;background:#7A8B99;border-radius:2px;margin-right:4px\"></span>Gastos generales (prorrateo): '+pr.toLocaleString('es-ES',{minimumFractionDigits:2})+' \u20ac</div>';legendHtml+='</div>';document.getElementById('resumen-detail-legend').innerHTML=legendHtml;var margenColor=p.margen>=0?'#27ae60':'#e74c3c';var certStr=p.has_cert?p.certificacion.toLocaleString('es-ES',{minimumFractionDigits:2})+' \u20ac':'Sin certificacion';var html2='<strong>CERTIFICADO</strong> '+certStr+' &mdash; <strong>COSTE TOTAL</strong> '+p.total_coste.toLocaleString('es-ES',{minimumFractionDigits:2})+' \u20ac';html2+='<br><span style=\"color:'+margenColor+';font-size:1.3rem;font-weight:700\">'+(p.margen>=0?'+':'')+p.margen.toLocaleString('es-ES',{minimumFractionDigits:2})+' \u20ac</span>';html2+='<br><span style=\"color:'+margenColor+'\">('+p.margen_pct+String.fromCharCode(37)+')</span>';document.getElementById('resumen-detail-summary').innerHTML=html2;}")
# 1) Global donut: overall cost breakdown
donut_js = "new Chart(document.getElementById('chartGlobalDonut'),{type:'doughnut',data:{"
donut_js += "labels:['Gastos Directos','Prorrateo GG+VEH','Mano de Obra'],"
donut_js += "datasets:[{data:[" + str(round(sum_directos,2)) + "," + str(round(sum_prorrateo,2)) + "," + str(round(sum_mo,2)) + "],"
donut_js += "backgroundColor:['#5BA3C9','#D4742C','#2B4C6F'],borderWidth:3,borderColor:'#F5F0E8'}]},"
donut_js += "options:{responsive:false,cutout:'60%',plugins:{legend:{position:'bottom',labels:{font:{size:13},padding:16}}}}});"
lines.append(donut_js)

# 2) Stacked horizontal bar: cost composition per project (certified only)
cert_projects = [p for p in proyectos_data if p['has_cert']]
cost_labels = [p['nombre'][:25] for p in cert_projects]
cost_direct = [round(p['gastos_directos'],2) for p in cert_projects]
cost_prorr = [round(p['prorrateo'],2) for p in cert_projects]
cost_mo = [round(p['mano_obra_coste'],2) for p in cert_projects]
lines.append("new Chart(document.getElementById('chartCostStack'),{type:'bar',data:{labels:" + json.dumps(cost_labels, ensure_ascii=False) + ",datasets:[{label:'Facturas directas',data:" + json.dumps(cost_direct) + ",backgroundColor:'#5BA3C9'},{label:'Prorrateo GG+VEH',data:" + json.dumps(cost_prorr) + ",backgroundColor:'#D4742C'},{label:'Mano de Obra',data:" + json.dumps(cost_mo) + ",backgroundColor:'#2B4C6F'}]},options:{indexAxis:'y',responsive:true,maintainAspectRatio:false,scales:{x:{stacked:true,ticks:{callback:function(v){return v.toLocaleString('es-ES')}}},y:{stacked:true}},plugins:{legend:{position:'top'}}}});")

# 3) Margen bar chart (green/red per project)
margen_proj = [p for p in projects_json_data if p['has_cert']]
margen_labels2 = [p['nombre'][:25] for p in margen_proj]
margen_vals2 = [p['margen'] for p in margen_proj]
margen_colors2 = ['#27ae60' if v >= 0 else '#e74c3c' for v in margen_vals2]
lines.append("new Chart(document.getElementById('chartMargenBar'),{type:'bar',data:{labels:" + json.dumps(margen_labels2, ensure_ascii=False) + ",datasets:[{label:'Margen (EUR)',data:" + json.dumps(margen_vals2) + ",backgroundColor:" + json.dumps(margen_colors2) + "}]},options:{indexAxis:'y',responsive:true,maintainAspectRatio:false,plugins:{legend:{display:false}},scales:{x:{ticks:{callback:function(v){return v.toLocaleString('es-ES')}}}},onClick:function(e,els){if(els.length>0){showResumenDetail(els[0].index);}}}});")

# 4) Comparison chart: cert vs cost (improved colors)
lines.append("new Chart(document.getElementById('chartComp'),{type:'bar',data:{labels:labels,datasets:[{label:'Certificacion',data:certData,backgroundColor:'#2B4C6F'},{label:'Coste Total',data:certData.map((c,i)=>directData[i]+prorrData[i]+moData[i]),backgroundColor:'#D4742C'}]},options:{indexAxis:'y',responsive:true,maintainAspectRatio:false,plugins:{legend:{position:'top'}}}});")
lines.append("new Chart(document.getElementById('chartProrrBar'),{type:'bar',data:{labels:labels,datasets:[{label:'Prorrateo (EUR)',data:prorrData,backgroundColor:'#D4742C'}]},options:{indexAxis:'y',responsive:true,maintainAspectRatio:false,plugins:{legend:{display:false}}}});")
lines.append("new Chart(document.getElementById('chartProrrPie'),{type:'pie',data:{labels:labels,datasets:[{data:prorrData,backgroundColor:['#2B4C6F','#D4742C','#5BA3C9','#7A8B99','#3B6A94','#4A7FB5','#8FA5B3','#C4A882','#2B3A4E','#6B8DA8']}]},options:{responsive:true,maintainAspectRatio:false,plugins:{legend:{position:'right',labels:{font:{size:10}}}}}});")
lines.append("new Chart(document.getElementById('chartGG'),{type:'doughnut',data:{labels:ggCatLabels,datasets:[{data:ggCatValues,backgroundColor:['#2B4C6F','#D4742C','#5BA3C9','#7A8B99','#3B6A94','#4A7FB5','#8FA5B3','#C4A882','#2B3A4E']}]},options:{responsive:true,maintainAspectRatio:false,plugins:{legend:{position:'right'}}}});")
lines.append("new Chart(document.getElementById('chartVeh'),{type:'doughnut',data:{labels:vehCatLabels,datasets:[{data:vehCatValues,backgroundColor:['#D4742C','#2B4C6F','#5BA3C9','#7A8B99','#3B6A94','#4A7FB5','#8FA5B3']}]},options:{responsive:true,maintainAspectRatio:false,plugins:{legend:{position:'right'}}}});")
# JSON data for chart click filtering
lines.append("var topProvFull=%s;" % json.dumps(top10_sup_full, ensure_ascii=False))
lines.append("var topProjFull=%s;" % json.dumps(top10_proj_full, ensure_ascii=False))

# Filter function for chart clicks
filter_js = []
filter_js.append('function filterByChart(type,idx){')
filter_js.append('  var val=type==\'prov\'?topProvFull[idx]:topProjFull[idx];')
filter_js.append('  var table=document.getElementById(\'facturasObraTable\');')
filter_js.append('  var rows=table.querySelectorAll(\'tbody tr\');')
filter_js.append('  var colIdx=type==\'prov\'?4:0;')
filter_js.append('  var shown=0;')
filter_js.append('  rows.forEach(function(row){')
filter_js.append('    var cells=row.querySelectorAll(\'td\');')
filter_js.append('    if(cells.length>colIdx){')
filter_js.append('      var txt=cells[colIdx].textContent.toLowerCase();')
filter_js.append('      if(txt.indexOf(val.toLowerCase())>=0){row.style.display=\'\';shown++;}')
filter_js.append('      else{row.style.display=\'none\';}')
filter_js.append('    }')
filter_js.append('  });')
filter_js.append('  var badge=document.getElementById(\'activeFilterBadge\');')
filter_js.append('  badge.style.display=\'inline\';badge.textContent=type==\'prov\'?\'Proveedor: \'+val:\'Proyecto: \'+val;')
filter_js.append('  document.getElementById(\'clearFilterBtn\').style.display=\'inline\';')
filter_js.append('  document.getElementById(\'clearFilterBtn\').scrollIntoView({behavior:\'smooth\',block:\'center\'});')
filter_js.append('}')
filter_js.append('function clearChartFilter(){')
filter_js.append('  var rows=document.getElementById(\'facturasObraTable\').querySelectorAll(\'tbody tr\');')
filter_js.append('  rows.forEach(function(row){row.style.display=\'\';});')
filter_js.append('  document.getElementById(\'activeFilterBadge\').style.display=\'none\';')
filter_js.append('  document.getElementById(\'clearFilterBtn\').style.display=\'none\';')
filter_js.append('  document.querySelector(\'.search-box\').value=\'\';')
filter_js.append('}')
for fl in filter_js:
    lines.append(fl)

# Top 10 suppliers chart (with onClick)
lines.append("var topProvChart=new Chart(document.getElementById('chartTopProv'),{type:'bar',data:{labels:%s,datasets:[{label:'Importe Total (EUR)',data:%s,backgroundColor:['#2B4C6F','#D4742C','#5BA3C9','#7A8B99','#3B6A94','#4A7FB5','#8FA5B3','#C4A882','#2B3A4E','#6B8DA8']}]},options:{indexAxis:'y',responsive:true,maintainAspectRatio:false,plugins:{legend:{display:false}},onClick:function(e,els){if(els.length>0){filterByChart('prov',els[0].index);}},scales:{x:{ticks:{callback:function(v){return v.toLocaleString('es-ES')}}}}}});" % (json.dumps(top10_labels, ensure_ascii=False), json.dumps(top10_values)))
# Top 10 projects chart (with onClick)
lines.append("var topProjChart=new Chart(document.getElementById('chartTopProj'),{type:'bar',data:{labels:%s,datasets:[{label:'Gastos Directos (EUR)',data:%s,backgroundColor:['#D4742C','#2B4C6F','#5BA3C9','#7A8B99','#3B6A94','#4A7FB5','#8FA5B3','#C4A882','#2B3A4E','#6B8DA8']}]},options:{indexAxis:'y',responsive:true,maintainAspectRatio:false,plugins:{legend:{display:false}},onClick:function(e,els){if(els.length>0){filterByChart('proj',els[0].index);}},scales:{x:{ticks:{callback:function(v){return v.toLocaleString('es-ES')}}}}}});" % (json.dumps(top10_proj_labels, ensure_ascii=False), json.dumps(top10_proj_values)))
lines.append('function generatePDF(){')
lines.append('var btn=document.getElementById(\'btnPDF\');')
lines.append('btn.innerHTML=\'&#9203; Generando PDF...\';')
lines.append('btn.disabled=true;btn.style.opacity=\'0.7\';')
lines.append('setTimeout(function(){')
lines.append('try{')
lines.append('var jsPDF=window.jspdf.jsPDF;')
lines.append('var doc=new jsPDF({unit:\'mm\',format:\'a4\',orientation:\'portrait\'});')
lines.append('var W=210,H=297,mL=15,mR=15,mT=15,mB=15;')
lines.append('var cW=W-mL-mR;')
lines.append('var y=mT;')
lines.append('var allProjects=typeof projectData!==\'undefined\'?projectData:[];')
lines.append('var sel=document.getElementById(\'pdfProjectSelect\');')
lines.append('var filterName=sel?sel.value:\'\';')
lines.append('var projects=filterName?allProjects.filter(function(p){return p.nombre===filterName;}):allProjects;')
lines.append('if(filterName&&projects.length===0){projects=allProjects;}')
lines.append('var d=new Date();')
lines.append('var ds=d.getDate()+\'/\'+(d.getMonth()+1)+\'/\'+d.getFullYear();')
lines.append('var titleSuffix=filterName?\' - \'+filterName:\'\';')
lines.append('function fmtE(v){')
lines.append('var neg=v<0;')
lines.append('var abs=Math.abs(v);')
lines.append('var intPart=Math.floor(abs);')
lines.append('var decPart=Math.round((abs-intPart)*100);')
lines.append('var decStr=decPart<10?\'0\'+decPart:String(decPart);')
lines.append('var s=String(intPart);')
lines.append('var result=\'\';')
lines.append('var count=0;')
lines.append('for(var i=s.length-1;i>=0;i--){')
lines.append('count++;')
lines.append('result=s[i]+result;')
lines.append('if(count%3===0&&i!==0) result=\'.\'+result;')
lines.append('}')
lines.append('return (neg?\'- \':\'\')+result+\',\'+decStr;')
lines.append('}')
lines.append('function checkPage(needed){')
lines.append('if(y+needed>H-mB){doc.addPage();y=mT;return true;}')
lines.append('return false;')
lines.append('}')
lines.append('function addLine(color){')
lines.append('doc.setDrawColor(color||212);doc.setLineWidth(0.5);')
lines.append('doc.line(mL,y,W-mR,y);y+=4;')
lines.append('}')
lines.append('var tC=0,tD=0,tP=0,tM=0,tO=0,tG=0,cc=0,nc=0;')
lines.append('projects.forEach(function(p){')
lines.append('tC+=p.certificacion;tD+=p.gastos_directos;tP+=p.prorrateo;')
lines.append('tM+=p.mano_obra;tO+=p.total_coste;tG+=p.margen;')
lines.append('if(p.has_cert)cc++;else nc++;')
lines.append('});')
lines.append('var mp=tC>0?(tG/tC*100).toFixed(1):\'0\';')
lines.append('doc.setFillColor(43,76,111);doc.rect(0,0,W,45,\'F\');')
lines.append('doc.setTextColor(255,255,255);doc.setFontSize(28);doc.setFont(\'helvetica\',\'bold\');')
lines.append('doc.text(\'ECO STRUCT\',W/2,20,{align:\'center\'});')
lines.append('doc.setFontSize(13);doc.setFont(\'helvetica\',\'normal\');')
lines.append('doc.text(\'Constructive Ecosen Spain 2.3\',W/2,28,{align:\'center\'});')
lines.append('doc.setFontSize(10);')
lines.append('doc.text(\'Dashboard Financiero\',W/2,35,{align:\'center\'});')
lines.append('y=60;')
lines.append('doc.setTextColor(43,58,78);doc.setFontSize(22);doc.setFont(\'helvetica\',\'bold\');')
lines.append('doc.text(\'INFORME FINANCIERO\',W/2,y,{align:\'center\'});y+=10;')
lines.append('doc.setFontSize(12);doc.setFont(\'helvetica\',\'normal\');')
lines.append('doc.setTextColor(90,106,122);')
lines.append('doc.text(\'Auditoria de Costes y Certificaciones\'+titleSuffix,W/2,y,{align:\'center\'});y+=8;')
lines.append('doc.setFontSize(11);')
lines.append('doc.text(\'Fecha: \'+ds,W/2,y,{align:\'center\'});y+=20;')
lines.append('addLine(212);y+=10;')
lines.append('doc.addPage();y=mT;')
lines.append('doc.setFontSize(14);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(212,116,44);')
lines.append('doc.text(\'1. RESUMEN EJECUTIVO\'+titleSuffix,mL,y);y+=3;addLine(212);y+=4;')
lines.append('var kpis=[')
lines.append('{l:\'Certificaciones\',v:fmtE(tC)+\' EUR\',bc:[43,76,111]},')
lines.append('{l:\'Gastos Directos\',v:fmtE(tD)+\' EUR\',bc:[233,69,96]},')
lines.append('{l:\'Gastos Comunes\',v:fmtE(tP)+\' EUR\',bc:[243,156,18]},')
lines.append('{l:\'Mano de Obra\',v:fmtE(tM)+\' EUR\',bc:[142,68,173]},')
lines.append('{l:\'MARGEN TOTAL\',v:fmtE(tG)+\' EUR\',bc:tG>=0?[46,204,113]:[231,76,60]}')
lines.append('];')
lines.append('var bx=mL,bw=(cW-8)/5;')
lines.append('kpis.forEach(function(k,i){')
lines.append('doc.setFillColor(245,240,232);doc.roundedRect(bx,y,bw,18,2,2,\'F\');')
lines.append('doc.setFillColor(k.bc[0],k.bc[1],k.bc[2]);doc.rect(bx,y,2,18,\'F\');')
lines.append('doc.setFontSize(6);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(90,106,122);')
lines.append('doc.text(k.l.toUpperCase(),bx+5,y+5);')
lines.append('doc.setFontSize(10);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(43,58,78);')
lines.append('doc.text(k.v,bx+5,y+12);')
lines.append('bx+=bw+2;')
lines.append('});')
lines.append('y+=24;')
lines.append('doc.setFontSize(11);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(43,76,111);')
lines.append('doc.text(\'Composicion de Costes Totales\',mL,y);y+=6;')
lines.append('doc.setFillColor(43,76,111);doc.rect(mL,y,cW,7,\'F\');')
lines.append('doc.setFontSize(8);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(255,255,255);')
lines.append('doc.text(\'Concepto\',mL+3,y+5);')
lines.append('doc.text(\'Importe\',mL+cW-50,y+5,{align:\'right\'});')
lines.append('doc.text(\'% del Total\',mL+cW-3,y+5,{align:\'right\'});')
lines.append('y+=7;')
lines.append('var items=[{n:\'Gastos Directos\',v:tD},{n:\'Prorrateo GG+VEH\',v:tP},{n:\'Mano de Obra\',v:tM}];')
lines.append('items.forEach(function(it,i){')
lines.append('if(i%2===0){doc.setFillColor(250,247,242);doc.rect(mL,y,cW,6,\'F\');}')
lines.append('doc.setFontSize(8);doc.setFont(\'helvetica\',\'normal\');doc.setTextColor(43,58,78);')
lines.append('doc.text(it.n,mL+3,y+4.5);')
lines.append('doc.text(fmtE(it.v)+\' EUR\',mL+cW-50,y+4.5,{align:\'right\'});')
lines.append('var pct=tO>0?(it.v/tO*100).toFixed(1):\'0\';')
lines.append('doc.text(pct+\'%\',mL+cW-3,y+4.5,{align:\'right\'});')
lines.append('y+=6;')
lines.append('});')
lines.append('doc.setFillColor(245,240,232);doc.rect(mL,y,cW,7,\'F\');')
lines.append('doc.setDrawColor(43,76,111);doc.setLineWidth(0.3);doc.line(mL,y,W-mR,y);')
lines.append('doc.setFontSize(8);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(43,58,78);')
lines.append('doc.text(\'TOTAL COSTES\',mL+3,y+5);')
lines.append('doc.text(fmtE(tO)+\' EUR\',mL+cW-50,y+5,{align:\'right\'});')
lines.append('doc.text(\'100%\',mL+cW-3,y+5,{align:\'right\'});')
lines.append('y+=12;')
lines.append('doc.setFontSize(9);doc.setTextColor(90,106,122);')
lines.append('doc.text(\'Eficiencia Coste/Certif: \'+(tC>0?(tO/tC*100).toFixed(1):\'0\')+\'% | Obras con beneficio: \'+cc+\' | Obras con perdida: \'+nc,mL,y);')
lines.append('if(!filterName){')
lines.append('doc.addPage();y=mT;')
lines.append('doc.setFontSize(14);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(212,116,44);')
lines.append('doc.text(\'2. DETALLE POR PROYECTO\',mL,y);y+=3;addLine(212);y+=4;')
lines.append('doc.setFillColor(43,76,111);doc.rect(mL,y,cW,7,\'F\');')
lines.append('doc.setFontSize(7);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(255,255,255);')
lines.append('var th=[\'#\',\'Proyecto\',\'Certif.\',\'G.Directos\',\'Prorrateo\',\'Mano Obra\',\'TOTAL\',\'MARGEN\'];')
lines.append('var tw=[8,52,28,28,28,28,28,28];')
lines.append('var tx=mL;')
lines.append('th.forEach(function(h,i){doc.text(h,tx+2,y+5);tx+=tw[i];});')
lines.append('y+=7;')
lines.append('projects.forEach(function(p,i){')
lines.append('checkPage(6);')
lines.append('if(i%2===0){doc.setFillColor(250,247,242);doc.rect(mL,y,cW,5.5,\'F\');}')
lines.append('var mc=p.margen>=0?[46,204,113]:[231,76,60];')
lines.append('doc.setFontSize(7);doc.setFont(\'helvetica\',\'normal\');doc.setTextColor(43,58,78);')
lines.append('var cv=p.has_cert?fmtE(p.certificacion):\'-\';')
lines.append('var vals=[String(i+1),p.nombre.length>28?p.nombre.substring(0,26)+\'..\':p.nombre,cv,fmtE(p.gastos_directos),fmtE(p.prorrateo),fmtE(p.mano_obra),fmtE(p.total_coste),fmtE(p.margen)];')
lines.append('tx=mL;')
lines.append('vals.forEach(function(v,j){')
lines.append('if(j===7){doc.setTextColor(mc[0],mc[1],mc[2]);doc.setFont(\'helvetica\',\'bold\');}')
lines.append('doc.text(v,tx+2,y+4);')
lines.append('if(j===7){doc.setTextColor(43,58,78);doc.setFont(\'helvetica\',\'normal\');}')
lines.append('tx+=tw[j];')
lines.append('});')
lines.append('y+=5.5;')
lines.append('});')
lines.append('checkPage(7);')
lines.append('doc.setFillColor(245,240,232);doc.rect(mL,y,cW,7,\'F\');')
lines.append('doc.setDrawColor(43,76,111);doc.setLineWidth(0.3);doc.line(mL,y,W-mR,y);')
lines.append('doc.setFontSize(7);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(43,58,78);')
lines.append('var tv=[\'\',\'TOTAL\',fmtE(tC),fmtE(tD),fmtE(tP),fmtE(tM),fmtE(tO),fmtE(tG)];')
lines.append('tx=mL;')
lines.append('tv.forEach(function(v,j){doc.text(v,tx+2,y+5);tx+=tw[j];});')
lines.append('var sectionNum=3;')
lines.append('}else{')
lines.append('var sectionNum=2;')
lines.append('}')
lines.append('doc.addPage();y=mT;')
lines.append('doc.setFontSize(14);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(212,116,44);')
lines.append('doc.text(sectionNum+\'. ANALISIS POR PROYECTO\'+titleSuffix,mL,y);y+=3;addLine(212);y+=4;')
lines.append('projects.forEach(function(p,i){')
lines.append('checkPage(35);')
lines.append('var mc=p.margen>=0?[46,204,113]:[231,76,60];')
lines.append('doc.setFontSize(11);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(43,76,111);')
lines.append('doc.text((i+1)+\'. \'+p.nombre,mL,y);y+=6;')
lines.append('var ky=y;var kbx=mL;var kbw=(cW-4)/3;')
lines.append('var kItems=[')
lines.append('{l:\'Certificacion\',v:p.has_cert?fmtE(p.certificacion)+\' EUR\':\'Sin certif.\',bc:[43,76,111]},')
lines.append('{l:\'Coste Total\',v:fmtE(p.total_coste)+\' EUR\',bc:[233,69,96]},')
lines.append('{l:\'Margen\',v:fmtE(p.margen)+\' EUR (\'+p.margen_pct+\'%)\',bc:mc}')
lines.append('];')
lines.append('kItems.forEach(function(k){')
lines.append('doc.setFillColor(245,240,232);doc.roundedRect(kbx,ky,kbw,14,1,1,\'F\');')
lines.append('doc.setFillColor(k.bc[0],k.bc[1],k.bc[2]);doc.rect(kbx,ky,1.5,14,\'F\');')
lines.append('doc.setFontSize(5.5);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(90,106,122);')
lines.append('doc.text(k.l.toUpperCase(),kbx+4,ky+5);')
lines.append('doc.setFontSize(9);doc.setFont(\'helvetica\',\'bold\');doc.setTextColor(43,58,78);')
lines.append('doc.text(k.v,kbx+4,ky+11);')
lines.append('kbx+=kbw+2;')
lines.append('});')
lines.append('y+=18;')
lines.append('doc.setFontSize(7);doc.setFont(\'helvetica\',\'bold\');')
lines.append('doc.setFillColor(43,76,111);doc.rect(mL,y,cW,5,\'F\');')
lines.append('doc.setTextColor(255,255,255);')
lines.append('doc.text(\'Componente\',mL+3,y+3.5);')
lines.append('doc.text(\'Importe\',mL+cW-45,y+3.5,{align:\'right\'});')
lines.append('doc.text(\'% Coste\',mL+cW-3,y+3.5,{align:\'right\'});')
lines.append('y+=5;')
lines.append('var comps=[{n:\'Gastos Directos\',v:p.gastos_directos},{n:\'Prorrateo GG+VEH\',v:p.prorrateo},{n:\'Mano de Obra\',v:p.mano_obra}];')
lines.append('comps.forEach(function(c2,j){')
lines.append('if(j%2===0){doc.setFillColor(250,247,242);doc.rect(mL,y,cW,4.5,\'F\');}')
lines.append('doc.setFontSize(7);doc.setFont(\'helvetica\',\'normal\');doc.setTextColor(43,58,78);')
lines.append('doc.text(c2.n,mL+3,y+3.2);')
lines.append('doc.text(fmtE(c2.v)+\' EUR\',mL+cW-45,y+3.2,{align:\'right\'});')
lines.append('var pct=p.total_coste>0?(c2.v/p.total_coste*100).toFixed(1):\'0\';')
lines.append('doc.text(pct+\'%\',mL+cW-3,y+3.2,{align:\'right\'});')
lines.append('y+=4.5;')
lines.append('});')
lines.append('doc.setFillColor(245,240,232);doc.rect(mL,y,cW,5,\'F\');')
lines.append('doc.setDrawColor(43,76,111);doc.setLineWidth(0.2);doc.line(mL,y,W-mR,y);')
lines.append('doc.setFontSize(7);doc.setFont(\'helvetica\',\'bold\');')
lines.append('doc.text(\'TOTAL\',mL+3,y+3.5);')
lines.append('doc.text(fmtE(p.total_coste)+\' EUR\',mL+cW-45,y+3.5,{align:\'right\'});')
lines.append('doc.text(\'100%\',mL+cW-3,y+3.5,{align:\'right\'});')
lines.append('y+=10;')
lines.append('});')
lines.append('var totalPages=doc.internal.getNumberOfPages();')
lines.append('for(var pg=1;pg<=totalPages;pg++){')
lines.append('doc.setPage(pg);')
lines.append('doc.setFontSize(6);doc.setFont(\'helvetica\',\'normal\');doc.setTextColor(170,170,170);')
lines.append('var footerText=\'ECO STRUCT - Constructive Ecosen Spain 2.3 | Informe generado el \'+ds;')
lines.append('if(filterName) footerText+=\' | Obra: \'+filterName;')
lines.append('footerText+=\' | Pagina \'+pg+\' de \'+totalPages;')
lines.append('doc.text(footerText,mL,H-8);')
lines.append('}')
lines.append('var fileName=\'ECO_STRUCT\';')
lines.append('if(filterName) fileName+=\'_obra-\'+filterName.substring(0,20);')
lines.append('fileName+=\'_Informe_\'+ds.split(\'/\').join(\'-\')+\'.pdf\';')
lines.append('doc.save(fileName);')
lines.append('btn.innerHTML=\'<svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/><polyline points="10 9 9 9 8 9"/></svg> Exportar PDF\';')
lines.append('btn.disabled=false;btn.style.opacity=\'1\';')
lines.append('}catch(err){')
lines.append('console.error(\'PDF Error:\',err);')
lines.append('alert(\'Error generando PDF: \'+err.message);')
lines.append('btn.innerHTML=\'Exportar PDF\';')
lines.append('btn.disabled=false;btn.style.opacity=\'1\';')
lines.append('}')
lines.append('},100);')
lines.append('}')
lines.append('function generateObraPDF(nombre){')
lines.append('var sel=document.getElementById(\'pdfProjectSelect\');')
lines.append('if(sel){sel.value=nombre;}')
lines.append('generatePDF();')
lines.append('}')

lines.append('function pdfFromDetail(){')
lines.append("  var el=document.getElementById('detail-title');")
lines.append("  var nombre=el.getAttribute('data-nombre');")
lines.append('  if(nombre) generateObraPDF(nombre);')
lines.append('}')

lines.append('</script>')
lines.append('</body>')
lines.append('</html>')

html = '\n'.join(lines)

with open('dashboard.html', 'w', encoding='utf-8') as f:
    f.write(html)

print("\ndashboard.html regenerado con datos CORRECTOS")
print("  Certificaciones: %s EUR" % fmt(sum_cert))
print("  Gastos Directos: %s EUR" % fmt(sum_directos))
print("  Gastos Comunes:  %s EUR" % fmt(total_gastos_comunes))
print("  Mano de Obra:    %s EUR" % fmt(sum_mo))
print("  MARGEN:          %s EUR (%.1f%%)" % (fmt(sum_margen), sum_margen/sum_cert*100 if sum_cert > 0 else 0))
